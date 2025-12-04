from openai import AsyncOpenAI
import json
import tempfile
import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import re


client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))



def generate_cv_filename(first_name, last_name, ext="docx"):
    # Normalize
    fn = (first_name or "").strip().lower()
    ln = (last_name or "").strip().lower()

    # Replace spaces with underscores
    fn = re.sub(r"\W+", "_", fn)
    ln = re.sub(r"\W+", "_", ln)

    # Random 6-digit number
    rand = random.randint(100000, 999999)

    return f"{fn}_{ln}_{rand}.{ext}"


def build_docx_from_json(cv_json):
    doc = Document()

    # Name header
    title = doc.add_heading(level=0)
    run = title.add_run(f"{cv_json['first_name']} {cv_json['last_name']}")
    run.bold = True
    run.font.size = Pt(20)

    # Contact line
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact.add_run(
        f"{cv_json.get('email','')} | {cv_json.get('phone','')}"
    ).bold = True

    doc.add_paragraph()

    # Summary
    doc.add_heading("Profile Summary", level=1)
    doc.add_paragraph(cv_json.get("summary", ""))

    # Skills
    doc.add_heading("Skills", level=1)
    skills_str = ", ".join(cv_json.get("skills", []))
    doc.add_paragraph(skills_str)

    # Experience
    doc.add_heading("Professional Experience", level=1)
    for exp in cv_json.get("experience", []):
        hdr = doc.add_paragraph()
        hdr.add_run(f"{exp.get('title','')} - {exp.get('company','')}").bold = True
        hdr.add_run(f" | {exp.get('location','')}")
        hdr.add_run(
            f" ({exp.get('start_date','')} - {exp.get('end_date','')})"
        )

        # Responsibilities
        if exp.get("responsibilities"):
            doc.add_paragraph().add_run("Responsibilities:").bold = True
            for r in exp["responsibilities"]:
                doc.add_paragraph(r, style="List Bullet")

        # Achievements
        if exp.get("achievements"):
            doc.add_paragraph().add_run("Key Achievements:").bold = True
            for a in exp["achievements"]:
                doc.add_paragraph(a, style="List Bullet")

    # Education
    doc.add_heading("Education", level=1)
    for edu in cv_json.get("education", []):
        p = doc.add_paragraph()
        p.add_run(f"{edu.get('degree','')} - {edu.get('institution','')}").bold = True
        p.add_run(f" ({edu.get('date','')})")

    return doc


async def generate_custom_cv(base_cv_text, job_text, user):

    prompt = f"""
    You are a world-class professional CV writer specialising in ATS optimisation, clarity, impact, and job alignment.
    
    Your task is to rewrite and restructure the candidate’s CV into a significantly improved version that:
    
    - Is noticeably stronger, clearer and more professional
    - Is fully ATS optimised with natural keyword usage
    - Corrects all grammar, style, clarity, and formatting issues
    - Embeds AT LEAST THREE high-value skills directly extracted from the job description
    - Highlights achievements over responsibilities
    - Uses strong action verbs and measurable outcomes wherever logically inferable
    - Modernises the structure to meet 2024–2025 hiring standards
    - Does NOT fabricate employment, education, tools, or certifications
    - May infer reasonable results ONLY when strongly implied by context
    - Keeps all factual information correct and anchored in the original CV
    
    Rewrite guidelines:
    - Transform weak responsibilities into strong bullet-point achievements
    - Remove filler or non-impactful language (“responsible for”, “helped with”, etc.)
    - Rewrite the summary into a sharp, value-driven pitch tailored to the job
    - Ensure skills are grouped logically (technical, soft, tools) and align with the job description
    - Maintain concise sentences; no long paragraphs
    - Standardise formatting, dates, job titles and location style
    
    Special requirements:
    - Ensure at least three job-description skills appear in the CV naturally
    - Improve readability and flow
    - Maintain JSON structure exactly with no deviation
    
    OUTPUT ONLY valid JSON.
    NO markdown.
    NO explanations.
    No commentary.
    
    JSON structure:
    {{
      "first_name": "",
      "last_name": "",
      "email": "",
      "phone": "",
      "address": "",
      "summary": "",
      "skills": [],
      "job_titles": [],
      "experience": [
        {{
          "title": "",
          "company": "",
          "location": "",
          "start_date": "",
          "end_date": "",
          "responsibilities": [],
          "achievements": []
        }}
      ],
      "education": [],
      "certifications": [],
      "languages": [],
      "additional_details": ""
    }}
    
    Base CV:
    {base_cv_text}
    
    Job Description:
    {job_text}
    
    Candidate Name:
    {user.get("first_name","")} {user.get("last_name","")}
    """

    response = await client.chat.completions.create(
        model="gpt-4.1",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.25
    )

    content = response.choices[0].message.content
    cleaned = content.replace("```json", "").replace("```", "").strip()

    cv_json = json.loads(cleaned)

    # Build DOCX
    doc = build_docx_from_json(cv_json)

    filename = generate_cv_filename(cv_json["first_name"], cv_json["last_name"])
    tmp_dir = tempfile.gettempdir()
    tmp_path = os.path.join(tmp_dir, filename)

    doc.save(tmp_path)

    return cv_json, tmp_path, filename
