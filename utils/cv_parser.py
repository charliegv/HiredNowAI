import pdfplumber
import docx
import os
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import subprocess
import tempfile


def extract_text_from_pdf(path):
    text = ""

    # 1. try normal text extraction
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception:
        pass

    if text.strip():
        return text

    # 2. fallback to OCR
    try:
        ocr_text = ""
        doc = fitz.open(path)
        for page in doc:
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text += pytesseract.image_to_string(img) + "\n"
        return ocr_text
    except Exception:
        return ""


def extract_text_from_docx(path):
    """Extract text from .docx using python-docx."""
    try:
        d = docx.Document(path)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception:
        return ""


def extract_text_from_doc(path):
    """
    Converts a legacy .doc file into .docx using LibreOffice headless,
    then extracts text normally using python-docx.
    """

    try:
        # Create temp dir
        tmp_dir = tempfile.mkdtemp()

        # Convert .doc → .docx
        subprocess.run([
            "soffice",                  # LibreOffice binary
            "--headless",
            "--convert-to", "docx",
            "--outdir", tmp_dir,
            path
        ], check=True)

        # Get the converted file name
        base = os.path.splitext(os.path.basename(path))[0]
        converted_path = os.path.join(tmp_dir, base + ".docx")

        if not os.path.exists(converted_path):
            print("Conversion failed: file not found:", converted_path)
            return ""

        # Extract text using python-docx
        import docx
        d = docx.Document(converted_path)
        return "\n".join(p.text for p in d.paragraphs)

    except Exception as e:
        print("DOC PARSE ERROR:", e)
        return ""



def extract_cv_text(path):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(path)

    if ext == ".docx":
        print("docx route")
        return extract_text_from_docx(path)

    if ext == ".doc":
        # try legacy .doc extractor
        text = extract_text_from_doc(path)
        if text.strip():
            return text
        # optional: fallback to docx logic if you ever convert .doc → .docx elsewhere
        return ""

    # fallback for plain text or unknown types
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""
